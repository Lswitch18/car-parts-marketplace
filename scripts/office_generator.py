#!/usr/bin/env python3
"""
Office Generator - GAID Platform
Gerador de Planilhas Excel (.xlsx) e Documentos Word (.docx) com suporte a exportação local e sincronização com Google Drive.
"""

import os
import sys
from datetime import datetime
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

OUTPUT_DIR = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "exports"))
os.makedirs(OUTPUT_DIR, exist_ok=True)

# ------------------------------------------------------------------------------
# 1. GERADOR DE PLANILHAS EXCEL (.XLSX)
# ------------------------------------------------------------------------------
def create_excel_spreadsheet(filename: str = "relatorio_financeiro.xlsx", data: list = None) -> str:
  """Cria uma planilha Excel (.xlsx) profissional e formatada."""
  filepath = os.path.join(OUTPUT_DIR, filename)
  wb = Workbook()
  ws = wb.active
  ws.title = "Relatório de Transações"

  # Ativar linhas de grade
  ws.views.sheetView[0].showGridLines = True

  # Estilos
  font_title = Font(name="Calibri", size=16, bold=True, color="FFFFFF")
  fill_title = PatternFill(
      start_color="1A1A1A", end_color="1A1A1A", fill_type="solid"
  )

  font_header = Font(name="Calibri", size=11, bold=True, color="FFFFFF")
  fill_header = PatternFill(
      start_color="2B4C7E", end_color="2B4C7E", fill_type="solid"
  )

  font_data = Font(name="Calibri", size=10, color="000000")
  font_total = Font(name="Calibri", size=11, bold=True, color="000000")

  border_thin = Side(border_style="thin", color="CCCCCC")
  box_border = Border(
      left=border_thin, right=border_thin, top=border_thin, bottom=border_thin
  )

  align_center = Alignment(horizontal="center", vertical="center")
  align_left = Alignment(horizontal="left", vertical="center")
  align_right = Alignment(horizontal="right", vertical="center")

  # Título
  ws.merge_cells("A1:G1")
  cell_title = ws["A1"]
  cell_title.value = "GAID - RELATÓRIO FINANCEIRO DE TRANSAÇÕES (JDM)"
  cell_title.font = font_title
  cell_title.fill = fill_title
  cell_title.alignment = align_center
  ws.row_dimensions[1].height = 40

  # Subtítulo Data
  ws.merge_cells("A2:G2")
  ws["A2"].value = (
      f"Gerado em: {datetime.now().strftime('%d/%m/%Y às %H:%M:%S')}"
  )
  ws["A2"].font = Font(name="Calibri", size=9, italic=True, color="555555")
  ws["A2"].alignment = Alignment(horizontal="left", vertical="center")

  # Cabeçalhos da Tabela
  headers = [
      "ID Transação",
      "Data",
      "Peça / Produto",
      "Status Pagamento",
      "Moeda",
      "Valor Bruto",
      "Taxa Líquida",
  ]
  ws.append([])  # Linha 3 vazia

  ws.row_dimensions[4].height = 28
  for col_num, header in enumerate(headers, 1):
    cell = ws.cell(row=4, column=col_num)
    cell.value = header
    cell.font = font_header
    cell.fill = fill_header
    cell.alignment = align_center
    cell.border = box_border

  # Dados Exemplo se não fornecidos
  if not data:
    data = [
        (
            "TX-849201",
            "2026-07-21",
            "Volante 無限 MF10 16 polegadas",
            "paid",
            "JPY",
            350000,
            315000,
        ),
        (
            "TX-849202",
            "2026-07-21",
            "Radiador Honda N-VAN Denso",
            "escrow",
            "JPY",
            12000,
            10800,
        ),
        (
            "TX-849203",
            "2026-07-21",
            "Amortecedor TEIN Flex Z Civic Type R",
            "pending_payment",
            "JPY",
            128000,
            115200,
        ),
        (
            "TX-849204",
            "2026-07-20",
            "Farol LED Nissan GT-R R35",
            "completed",
            "JPY",
            210000,
            189000,
        ),
    ]

  start_row = 5
  total_bruto = 0
  total_liquido = 0

  for row_idx, row_data in enumerate(data, start=start_row):
    ws.row_dimensions[row_idx].height = 22
    for col_idx, value in enumerate(row_data, start=1):
      cell = ws.cell(row=row_idx, column=col_idx)
      cell.value = value
      cell.font = font_data
      cell.border = box_border

      if col_idx in [1, 2, 4, 5]:
        cell.alignment = align_center
      elif col_idx in [6, 7]:
        cell.alignment = align_right
        cell.number_format = '"¥ "#,##0'

    total_bruto += row_data[5]
    total_liquido += row_data[6]

  # Linha de Totais
  total_row = start_row + len(data)
  ws.row_dimensions[total_row].height = 26

  ws.cell(row=total_row, column=1).value = "TOTAL CONSOLIDADO"
  ws.cell(row=total_row, column=1).font = font_total
  ws.cell(row=total_row, column=1).alignment = align_left

  for c in range(1, 6):
    ws.cell(row=total_row, column=c).border = box_border
    ws.cell(row=total_row, column=c).font = font_total

  c_bruto = ws.cell(row=total_row, column=6)
  c_bruto.value = total_bruto
  c_bruto.font = font_total
  c_bruto.alignment = align_right
  c_bruto.number_format = '"¥ "#,##0'
  c_bruto.border = box_border

  c_liq = ws.cell(row=total_row, column=7)
  c_liq.value = total_liquido
  c_liq.font = font_total
  c_liq.alignment = align_right
  c_liq.number_format = '"¥ "#,##0'
  c_liq.border = box_border

  # Ajuste de largura das colunas
  for col in ws.columns:
    max_len = max(len(str(cell.value or '')) for cell in col)
    col_letter = get_column_letter(col[0].column)
    ws.column_dimensions[col_letter].width = max(max_len + 4, 14)

  wb.save(filepath)
  print(f"✅ Planilha Excel gerada com sucesso em: {filepath}")
  return filepath


# ------------------------------------------------------------------------------
# 2. GERADOR DE DOCUMENTOS WORD (.DOCX)
# ------------------------------------------------------------------------------
def create_word_document(
    filename: str = "contrato_parceria_b2b.docx",
    title: str = "CONTRATO DE PARCERIA E SERVIÇOS B2B",
) -> str:
  """Cria um documento Word (.docx) profissional formatado."""
  filepath = os.path.join(OUTPUT_DIR, filename)
  doc = Document()

  # Configurar Margens
  sections = doc.sections
  for section in sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

  # Cabeçalho Principal
  p_header = doc.add_paragraph()
  p_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
  run_title = p_header.add_run(f"GAID PLATFORM - JDM MARKETPLACE\n{title}")
  run_title.font.name = "Arial"
  run_title.font.size = Pt(16)
  run_title.font.bold = True
  run_title.font.color.rgb = RGBColor(26, 26, 26)

  doc.add_paragraph().paragraph_format.space_after = Pt(12)

  # Subtítulo de identificação
  p_meta = doc.add_paragraph()
  p_meta.paragraph_format.space_after = Pt(18)
  run_meta = p_meta.add_run(
      f"Documento Gerado Oficialmente em: {datetime.now().strftime('%d/%m/%Y às %H:%M')}\nCódigo de Registro: CTR-B2B-2026-JDM-9401"
  )
  run_meta.font.name = "Arial"
  run_meta.font.size = Pt(10)
  run_meta.font.italic = True
  run_meta.font.color.rgb = RGBColor(100, 100, 100)

  # Seção 1: Das Partes
  doc.add_heading("1. Das Partes Contratantes", level=1)
  p1 = doc.add_paragraph(
      "De um lado, a plataforma GAID JDM Car Parts Marketplace (doravante"
      " denominada CONTRATADA), e de outro lado, a empresa parceira de"
      " logística e fornecimento de autopeças (doravante denominada"
      " CONTRATANTE), celebram o presente instrumento contratual sob as"
      " cláusulas abaixo estipuladas."
  )
  p1.paragraph_format.line_spacing = 1.15
  p1.paragraph_format.space_after = Pt(12)

  # Seção 2: Do Objeto e Serviços
  doc.add_heading("2. Do Objeto e Termos da Parceria", level=1)
  p2 = doc.add_paragraph(
      "O presente contrato tem como objeto a prestação de serviços de"
      " intermediação de vendas, custódia de pagamento via cartão/Konbini"
      " (Escrow) e sincronização logística no mercado automotivo do Japão."
  )
  p2.paragraph_format.line_spacing = 1.15
  p2.paragraph_format.space_after = Pt(12)

  # Tabela com especificações
  table = doc.add_table(rows=4, cols=3)
  table.alignment = WD_TABLE_ALIGNMENT.CENTER

  headers = ["Item / Serviço", "Periodicidade", "Valor (JPY)"]
  hdr_cells = table.rows[0].cells
  for idx, text in enumerate(headers):
    hdr_cells[idx].text = text
    hdr_cells[idx].paragraphs[0].runs[0].font.bold = True
    hdr_cells[idx].paragraphs[0].runs[0].font.color.rgb = RGBColor(
        255, 255, 255
    )

  services = [
      ("Integração WMS Logistix B2B", "Mensal", "¥ 60.000"),
      ("Manutenção de API Key & Endpoints", "Mensal", "¥ 15.000"),
      ("Custódia de Pagamento (Escrow)", "Por Transação", "10% Comissão"),
  ]

  for row_idx, data_tuple in enumerate(services, start=1):
    row_cells = table.rows[row_idx].cells
    for col_idx, cell_value in enumerate(data_tuple):
      row_cells[col_idx].text = cell_value

  doc.add_paragraph().paragraph_format.space_after = Pt(24)

  # Assinaturas
  doc.add_heading("3. Assinatura e Homologação Digital", level=1)
  p_sig = doc.add_paragraph()
  p_sig.paragraph_format.space_before = Pt(36)
  p_sig.add_run(
      "________________________________________          ________________________________________\n"
  )
  p_sig.add_run(
      "GAID JDM Marketplace                              Tech in Technologies"
      " (Contratante)\n"
  )
  p_sig.add_run(
      "Representante Legal                               Diretoria Executiva"
  )

  doc.save(filepath)
  print(f"✅ Documento Word gerado com sucesso em: {filepath}")
  return filepath


# ------------------------------------------------------------------------------
# 3. SCRIPT EXECUTION
# ------------------------------------------------------------------------------
if __name__ == "__main__":
  excel_path = create_excel_spreadsheet()
  word_path = create_word_document()

  print("\n🎉 Todos os arquivos Office foram criados com sucesso!")
  print(f"📄 Excel: {excel_path}")
  print(f"📝 Word:  {word_path}")
